"""Turn a hand-written Ren'Py scene into an emitter that regenerates it from
the author's docx.

    python convert_scene.py <docx> <patterns.json> <scene.rpy> <out_dir> [--check]

Writes <out_dir>/emit_<scene>.py. With --check it also RUNS the new emitter
to a temp file and diffs the result against the original scene, so the
conversion is verified against the one oracle that exists: the file itself.

WHY THIS EXISTS
---------------
Scenes converted before the emitters existed are correct today and drift
tomorrow: every document-wide edit the author makes has to be found and
re-applied by hand in each one, while the emitted scenes take it for free
on the next run. One maintenance model for the whole game is the point.

WHAT IT PRESERVES, AND HOW
--------------------------
A hand-written scene is prose from the docx interleaved with STAGING --
`scene`, `show`, `hide`, `with`, `play`, `nvl`, `$`, `call`, `jump`, `if`
-- and the staging is the part a human decided. So:

  * Every say line is matched to its docx paragraph and REPLACED by a
    lookup, so the prose comes from the document at emit time.
  * Everything else is kept VERBATIM, attached to the paragraph it precedes
    (`before`) or, after the last one, as the scene's tail.
  * The scene's header comments become the emitter's docstring, verbatim.
    They record staging reasoning ("the sword is REVEALED, not mentioned")
    that must not be lost in the move.

ANCHORS, RESOLVED IN ORDER. Every staging point is keyed to a short prefix
of the paragraph it precedes, and the emitter resolves them SEQUENTIALLY,
each search starting just past the previous match. So an anchor does not
need to be unique in the document -- only the FIRST occurrence after the
previous anchor has to be the right one, and that is what is verified
here. This matters: a bare ellipsis line, or a character saying only
"...", appears many times in a novel and can never be unique in full.

An edit that shifts the document cannot move a `show` onto the wrong
beat; an edit that rewrites an ANCHORED paragraph makes the emitter fail
loudly, naming the anchor, which is the correct outcome.

THE CASES A NAIVE VERSION GETS WRONG, all met on the first run:

  * The docx patterns must be loaded INTO script_diff's globals, not just
    returned. normalize() reads the globals; with the defaults it does not
    know this project's "To Claude" notes and nothing annotated matches.
  * Paragraphs are strip()ped, not rstrip()ped. A leading soft line break
    inside a paragraph defeats the ^-anchored speaker pattern otherwise.
  * A say whose speaker variable cannot be derived from the document
    (`centered "Ten hours later"`) carries the variable explicitly.
  * A say the scene TRANSFORMED -- a document line shown as a sized,
    centered card -- is matched with engine tags stripped, kept VERBATIM,
    and its paragraph is not emitted a second time.
  * A say inside a Ren'Py `if` branch keeps its deeper indent.
  * A say that appears OUT OF ORDER -- one scene repeats an earlier scene's
    opening lines as a flashback catching the frame -- is looked up
    absolutely and emitted where the scene put it.
  * Spec notes, choice cards and scenario headers in the document are
    skipped by script_diff.prose_mask(), the SAME state machine the diff
    uses, so the emitter and the diff can never disagree about them.
  * Trailing author annotations are stripped from the emitted text with
    script_diff's own ANNOTATION pattern.

A prose paragraph in range that the hand-written scene simply does not
contain is REPORTED, not silently emitted or silently dropped -- a human
decides whether the scene or the document is right.

THE ORACLE. --check compares non-blank lines after normalizing what the
two sides legitimately disagree on: the hand-written files use straight,
backslash-escaped quotes where the document has curly ones. Blank-line
spacing is not compared; content and order are, strictly.
"""
import io
import os
import re
import subprocess
import sys
import tempfile

# The report quotes the author's prose, which has curly quotes and dashes
# a Windows console cannot encode by default. Silence is not an option.
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import script_diff  # noqa: E402

_SAY = re.compile(r'^(\s*)([a-z_][a-z0-9_]*\s+)?"(.*)"\s*$')
_LABEL = re.compile(r"^label\s+([A-Za-z_][\w]*)\s*:")
_SPEAKER_PREFIX = re.compile(r"^([^:]{1,40}):\s")
_TAGS = re.compile(r"\{[^}]*\}")

ANCHOR_MIN = 40
DEFAULT_INDENT = "    "


def unescape(s):
    return (s.replace('\\"', '"').replace("[[", "[").replace("{{", "{")
             .replace("\\\\", "\\"))


def load_docx(path):
    import docx
    return [p.text.strip() for p in docx.Document(path).paragraphs]


def parse_scene(path):
    lines = io.open(path, encoding="utf-8").read().split("\n")
    for i, l in enumerate(lines):
        m = _LABEL.match(l)
        if m:
            header = lines[:i]
            while header and not header[-1].strip():
                header.pop()
            return header, m.group(1), lines[i + 1:]
    raise SystemExit("convert_scene: no `label` in %s" % path)


def anchor_for(paras, idx, lo):
    """Shortest prefix of paras[idx] whose FIRST occurrence at or after `lo`
    is idx. Later duplicates are fine; the emitter resolves in order."""
    text = paras[idx]
    n = min(ANCHOR_MIN, len(text))
    while True:
        pre = text[:n]
        first = next((k for k in range(lo, len(paras)) if pre in paras[k]), None)
        if first == idx:
            return pre
        if n >= len(text):
            # Identical in full to an earlier paragraph -- a bare ellipsis, a
            # character saying only "...". Anchor on the PRECEDING non-empty
            # paragraph as well; the pair is what the emitter's find() matches.
            prev = next((k for k in range(idx - 1, -1, -1) if paras[k]), None)
            if prev is None or prev < lo:
                raise SystemExit(
                    "convert_scene: paragraph %d cannot be anchored even with "
                    "context: %r" % (idx, text[:60]))
            ctx = anchor_for(paras, prev, lo)
            pair = (ctx, text)
            first = next((k for k in range(lo, len(paras))
                          if _pair_hit(paras, k, pair)), None)
            if first != idx:
                raise SystemExit(
                    "convert_scene: paragraph %d not unique even with context "
                    "%r: %r" % (idx, ctx[:40], text[:60]))
            return pair
        n += 12


def _pair_hit(paras, k, pair):
    """A context anchor hits at k when paras[k] contains the text AND the
    previous non-empty paragraph contains the context."""
    ctx, text = pair
    if text not in paras[k]:
        return False
    prev = next((j for j in range(k - 1, -1, -1) if paras[j]), None)
    return prev is not None and ctx in paras[prev]


def trim(ls):
    ls = list(ls)
    while ls and not ls[0].strip():
        ls.pop(0)
    while ls and not ls[-1].strip():
        ls.pop()
    return ls


def convert(docx_path, patterns_path, scene_path, out_dir, check=False):
    (script_diff.SPEAKER, script_diff.SPEC_START, script_diff.SPEC_LINE,
     script_diff.ANNOTATION) = script_diff.load_patterns(patterns_path)
    paras = load_docx(docx_path)
    normed = [script_diff.normalize(p) if p else "" for p in paras]
    prose = script_diff.prose_mask(paras)
    header, label, body = parse_scene(scene_path)

    # ---- pass 1: classify lines, match says to paragraphs ----
    # item: ("other", line) | ("say", line, indent, var, idx, in_range, literal)
    items, cursor = [], 0
    unmatched, out_of_order, transformed = [], [], []
    speakers = {}
    for line in body:
        m = _SAY.match(line)
        if not m:
            items.append(("other", line))
            continue
        indent, var, text = m.group(1), (m.group(2) or "").strip(), m.group(3)
        key = script_diff.normalize(unescape(text))
        literal = False
        hit = next((k for k in range(cursor, len(paras)) if normed[k] == key), None)
        if hit is None:
            key2 = script_diff.normalize(_TAGS.sub("", unescape(text)))
            hit = next((k for k in range(cursor, len(paras)) if normed[k] == key2), None)
            if hit is not None:
                literal = True
                transformed.append(line.strip()[:60])
        in_range = True
        if hit is None:
            hit = next((k for k in range(len(paras)) if normed[k] == key), None)
            in_range = False
            if hit is None:
                unmatched.append(line.strip()[:70])
                items.append(("other", line))
                continue
            out_of_order.append((hit, line.strip()[:50]))
        else:
            cursor = hit + 1
        if var and var != "centered":
            sp = _SPEAKER_PREFIX.match(paras[hit])
            if sp:
                speakers.setdefault(sp.group(1).strip(), var)
        items.append(("say", line, indent, var, hit, in_range, literal))

    says_in = [it for it in items if it[0] == "say" and it[5]]
    if not says_in:
        raise SystemExit("convert_scene: no say line in %s matched the docx" % scene_path)
    start, last = says_in[0][4], says_in[-1][4]
    end = next((k for k in range(last + 1, len(paras)) if paras[k]), len(paras))

    def derived_var(idx):
        sp = _SPEAKER_PREFIX.match(paras[idx])
        return speakers.get(sp.group(1).strip()) if sp else None

    # ---- pass 2: fold staging onto the paragraph it precedes ----
    stages = []
    preamble, pending = [], []
    seen_first, cur = False, None
    for it in items:
        if it[0] == "other":
            pending.append(it[1])
            continue
        _, line, indent, var, idx, in_range, literal = it
        if not in_range:
            # Repeated from elsewhere: attach after the previous in-range say.
            cur["after"].extend(trim(pending) + [("SAY", var, idx, indent)])
            pending = []
            continue
        st = {"idx": idx, "before": [], "after": [], "var": None,
              "literal": line if literal else None,
              "indent": indent if indent != DEFAULT_INDENT else None}
        if not literal and var != (derived_var(idx) or ""):
            st["var"] = var  # "" forces narration, a name forces that speaker
        if not seen_first:
            preamble, seen_first = trim(pending), True
        else:
            st["before"] = trim(pending)
        pending = []
        stages.append(st)
        cur = st
    tail = trim(pending)

    matched = {st["idx"] for st in stages}
    force = sorted(k for k in matched if not prose[k])
    dropped = [k for k in range(start, end)
               if paras[k] and prose[k] and k not in matched]

    # ---- pass 3: anchors, resolved in order ----
    start_anchor = anchor_for(paras, start, 0)
    lo = start
    kept = []
    for st in stages:
        if not (st["before"] or st["after"] or st["var"] is not None
                or st["literal"] or st["indent"] or st["idx"] in force):
            continue
        st["anchor"] = anchor_for(paras, st["idx"], lo)
        lo = st["idx"] + 1
        st["after"] = [("SAY", x[1], anchor_for(paras, x[2], 0), x[3])
                       if isinstance(x, tuple) else x for x in st["after"]]
        kept.append(st)
    end_anchor = anchor_for(paras, end, lo) if end < len(paras) else None
    force_anchors = [anchor_for(paras, k, start) for k in force]

    # ---- write the emitter ----
    name = os.path.splitext(os.path.basename(scene_path))[0]
    out_path = os.path.join(out_dir, "emit_%s.py" % name)
    tools_dir = os.path.dirname(os.path.abspath(script_diff.__file__))
    doc = "\n".join((l[2:] if l.startswith("# ") else l[1:]) if l.startswith("#") else l
                    for l in header)
    g = ['"""%s -- emitted from the author\'s docx.' % name, "",
         "GENERATED FROM THE HAND-WRITTEN SCENE by convert_scene.py, then verified",
         "against it line for line. The original file's header follows verbatim;",
         "it is the staging reasoning, and it is why this file looks the way it does.",
         "", "-" * 72, doc.replace('"""', "'''"), "-" * 72, "",
         "Never hand-edit the .rpy this writes. Re-run this instead.", '"""',
         "import io", "import os", "import re", "import sys", "",
         "sys.path.insert(0, %r)" % tools_dir,
         "import docx  # noqa: E402",
         "import script_diff  # noqa: E402", "",
         "DOCX = %r" % os.path.abspath(docx_path),
         "PATTERNS = %r" % os.path.abspath(patterns_path),
         "OUT = %r" % os.path.abspath(os.path.dirname(scene_path)), "",
         "SPEAKERS = {"]
    g += ["    %r: %r," % (k, v) for k, v in sorted(speakers.items())]
    g += ["}", "", HELPERS, "def main():",
          "    (script_diff.SPEAKER, script_diff.SPEC_START, script_diff.SPEC_LINE,",
          "     script_diff.ANNOTATION) = script_diff.load_patterns(PATTERNS)",
          "    paras = [p.text.strip() for p in docx.Document(DOCX).paragraphs]",
          "    prose = script_diff.prose_mask(paras)", "",
          "    start = find(paras, %r, 0)" % (start_anchor,),
          "    stages = ["]
    for st in kept:
        g.append("        {")
        g.append("            'anchor': %r," % (st["anchor"],))
        for key in ("var", "literal", "indent"):
            if st[key] is not None:
                g.append("            %r: %r," % (key, st[key]))
        for key in ("before", "after"):
            if st[key]:
                g.append("            %r: [" % key)
                for x in st[key]:
                    if isinstance(x, tuple):
                        g.append("                # OUT OF ORDER on purpose: repeated from elsewhere")
                        g.append("                # in the document. Looked up absolutely.")
                        g.append("                %r," % (x,))
                    else:
                        g.append("                %r," % x)
                g.append("            ],")
        g.append("        },")
    g += ["    ]",
          "    # Matched prose the spec mask would otherwise drop.",
          "    force = %r" % force_anchors,
          "    cur = start",
          "    resolved = {}",
          "    for st in stages:",
          "        cur = find(paras, st['anchor'], cur)",
          "        resolved[cur] = st",
          "        cur += 1"]
    if end_anchor:
        g.append("    end = find(paras, %r, cur)" % (end_anchor,))
    else:
        g += ["    # This scene runs to the END OF THE DOCUMENT. Capped, not trusted:",
              "    # when the author writes past it, fail loudly and ask for an anchor.",
              "    end = len(paras)",
              "    if sum(1 for k in range(start, end) if paras[k]) > %d:" % (len(says_in) + 40),
              "        raise SystemExit('emit_%s: the document has grown past this "
              "scene; give it an END ANCHOR')" % name]
    g += ["", "    out = ["]
    g += ["        %r," % l for l in header]
    # The banner every generated scene carries. It is what tells the next
    # session -- and the repo's own tooling -- not to hand-edit this file.
    g.append("        '#',")
    g.append("        %r," % ("# ⚠ GENERATED by tools/emit_%s.py from the author's docx. "
                              "Re-run it; never edit this file." % name))
    g.append("        '',")
    g.append("        'label %s:'," % label)
    g += ["        %r," % l for l in preamble]
    g += ["    ]",
          "    forced = {find(paras, a, start) for a in force}",
          "    out += block(paras, prose, forced, start, end, resolved)"]
    g += ["    out.append(%r)" % l for l in tail]
    g += ["    out.append('')",
          "    path = os.path.join(OUT, %r)" % (name + ".rpy"),
          "    io.open(path, 'w', encoding='utf-8', newline='\\n').write('\\n'.join(out))",
          "    print('  wrote %%-24s %%d lines' %% (%r, len(out)))" % (name + ".rpy"),
          "    return 0", "", "",
          "if __name__ == '__main__':", "    sys.exit(main())", ""]
    io.open(out_path, "w", encoding="utf-8", newline="\n").write("\n".join(g))

    rep = ["convert_scene: %s -> %s" % (os.path.basename(scene_path), os.path.basename(out_path)),
           "  paragraphs %d..%d, %d say lines matched, %d staging points"
           % (start, end, len(says_in), len(kept))]
    if transformed:
        rep.append("  TRANSFORMED (kept verbatim, paragraph not re-emitted): %d" % len(transformed))
        rep += ["    %s" % t for t in transformed]
    if out_of_order:
        rep.append("  OUT OF ORDER (absolute lookups): %d" % len(out_of_order))
        rep += ["    para %d  %s" % (k, t) for k, t in out_of_order]
    if force:
        rep.append("  FORCED past the spec mask (matched prose the mask would drop): %d" % len(force))
        rep += ["    para %d  %s" % (k, paras[k][:60]) for k in force]
    if unmatched:
        rep.append("  UNMATCHED say lines (kept verbatim -- CHECK THESE): %d" % len(unmatched))
        rep += ["    %s" % t for t in unmatched]
    if dropped:
        rep.append("  DROPPED prose in range (in the document, not in the scene -- CHECK): %d"
                   % len(dropped))
        rep += ["    para %d  %s" % (k, paras[k][:70]) for k in dropped]
    print("\n".join(rep))
    return verify(out_path, scene_path, name) if check else 0


HELPERS = '''
def find(paras, needle, frm):
    """A plain anchor is a prefix of the target paragraph. A PAIR anchor is
    (context, text): the target contains `text` and the previous non-empty
    paragraph contains `context` -- for lines that repeat verbatim."""
    for n in range(frm, len(paras)):
        if isinstance(needle, tuple):
            if needle[1] in paras[n]:
                prev = next((j for j in range(n - 1, -1, -1) if paras[j]), None)
                if prev is not None and needle[0] in paras[prev]:
                    return n
        elif needle in paras[n]:
            return n
    raise SystemExit("anchor not found in docx after paragraph %d: %r" % (frm, needle))


def split_speaker(t):
    m = re.match(r"^([^:]{1,40}):\\s*(.*)$", t)
    if not m or m.group(1).strip() not in SPEAKERS:
        return None, t
    return SPEAKERS[m.group(1).strip()], m.group(2).strip()


def esc(s):
    return s.replace("\\\\", "\\\\\\\\").replace('"', '\\\\"').replace("[", "[[") \\
            .replace("{", "{{")


def say(t, indent="    ", var=None):
    t = script_diff.ANNOTATION.sub("", t).strip()
    who, body = split_speaker(t)
    if var is not None:
        who = var or None
    body = body.strip()
    if len(body) >= 2 and body[0] in "\\u201c\\"" and body[-1] in "\\u201d\\"":
        body = body[1:-1]
    return indent + ('%s "%s"' % (who, esc(body)) if who else '"%s"' % esc(body))


def block(paras, prose, forced, a, b, resolved):
    out = []
    for k in range(a, b):
        t = paras[k]
        if not t or (not prose[k] and k not in forced):
            continue
        st = resolved.get(k, {})
        out += st.get("before", [])
        if st.get("literal") is not None:
            out.append(st["literal"])
        else:
            out.append(say(t, st.get("indent") or "    ", st.get("var")))
        out.append("")
        for x in st.get("after", []):
            if isinstance(x, tuple):
                out.append(say(paras[find(paras, x[2], 0)], x[3] or "    ", x[1]))
                out.append("")
            else:
                out.append(x)
    return out

'''


def _norm_lines(text):
    out = []
    for l in text.split("\n"):
        if l.lstrip().startswith("#"):
            continue  # comments are not content; the emitted file gains a banner
        l = l.replace('\\"', '"')
        for a, b in (("’", "'"), ("‘", "'"), ("“", '"'), ("”", '"')):
            l = l.replace(a, b)
        if l.strip():
            out.append(l.rstrip())
    return out


def verify(emitter, original, name):
    tmp = tempfile.mkdtemp(prefix="convert_scene_")
    src = io.open(emitter, encoding="utf-8").read()
    # A lambda, because re.sub interprets backslashes in a plain replacement
    # string and a Windows path is nothing but backslashes.
    src = re.sub(r"^OUT = .*$", lambda m: "OUT = %r" % tmp, src, count=1, flags=re.M)
    tmp_emitter = os.path.join(tmp, os.path.basename(emitter))
    io.open(tmp_emitter, "w", encoding="utf-8", newline="\n").write(src)
    r = subprocess.run([sys.executable, tmp_emitter], capture_output=True, text=True,
                       encoding="utf-8", errors="replace")
    if r.returncode:
        print("  CHECK: emitter FAILED to run:\n" + (r.stderr or r.stdout))
        return 2
    a = _norm_lines(io.open(original, encoding="utf-8").read())
    b = _norm_lines(io.open(os.path.join(tmp, name + ".rpy"), encoding="utf-8").read())
    if a == b:
        print("  CHECK: IDENTICAL to the original (%d non-blank lines)" % len(a))
        return 0
    import difflib
    print("  CHECK: DIFFERS from the original:")
    for l in difflib.unified_diff(a, b, "original", "emitted", lineterm="", n=1):
        print("    " + l)
    return 1


def main(argv):
    check = "--check" in argv
    argv = [x for x in argv if x != "--check"]
    if len(argv) != 4:
        sys.exit(__doc__.strip().split("\n")[3].strip())
    return convert(*argv, check=check)


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
